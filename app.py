import os
import io
import sys
import json
import sqlite3
import hashlib
import secrets
import uuid
import shutil
from functools import wraps
from flask import (
    Flask, render_template, request, jsonify, send_file,
    session, redirect, url_for, send_from_directory
)
from openai import OpenAI
import anthropic
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import re


# ---------- Path helpers for PyInstaller ----------
def _bundle_dir():
    """Read-only bundled resources (templates, static, tinymce)."""
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


def _app_dir():
    """Writable runtime directory (db, uploads, template.docx) — next to exe."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


# On first run from exe, copy writable files to app dir if not present
if getattr(sys, 'frozen', False):
    for _fname in ('template.docx', 'zhidu.db'):
        _bundled = os.path.join(_bundle_dir(), _fname)
        _runtime = os.path.join(_app_dir(), _fname)
        if os.path.exists(_bundled) and not os.path.exists(_runtime):
            shutil.copy2(_bundled, _runtime)


app = Flask(__name__,
            template_folder=os.path.join(_bundle_dir(), 'templates'),
            static_folder=os.path.join(_bundle_dir(), 'static'),
            static_url_path='/static')
app.secret_key = secrets.token_hex(32)

# Serve TinyMCE from node_modules
TINYMCE_PATH = os.path.join(_bundle_dir(), 'node_modules', 'tinymce')

# Upload folder for images
UPLOAD_FOLDER = os.path.join(_app_dir(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/tinymce/<path:filename>')
def tinymce_static(filename):
    return send_from_directory(TINYMCE_PATH, filename)

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

@app.route('/api/upload-image', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': 'Empty filename'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg'):
        return jsonify({'error': 'Unsupported image format'}), 400
    filename = uuid.uuid4().hex + ext
    f.save(os.path.join(UPLOAD_FOLDER, filename))
    return jsonify({'location': '/uploads/' + filename})

DB_PATH = os.path.join(_app_dir(), 'zhidu.db')

# ---------- Database ----------
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS admin_users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS ai_config (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        api_key TEXT NOT NULL,
        model TEXT NOT NULL,
        base_url TEXT NOT NULL,
        api_type TEXT NOT NULL DEFAULT 'openai',
        is_active INTEGER DEFAULT 0
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS site_config (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL
    )''')
    # Migrate: add api_type column if missing (for existing DBs)
    cols = [row[1] for row in c.execute('PRAGMA table_info(ai_config)').fetchall()]
    if 'api_type' not in cols:
        c.execute("ALTER TABLE ai_config ADD COLUMN api_type TEXT NOT NULL DEFAULT 'openai'")
    # Default site config
    c.execute("INSERT OR IGNORE INTO site_config (key, value) VALUES (?, ?)",
              ('drawio_url', 'https://embed.diagrams.net'))
    # Default admin: admin / admin123 (should change after first login)
    default_pw = hashlib.sha256('admin123'.encode()).hexdigest()
    c.execute('SELECT COUNT(*) FROM admin_users')
    if c.fetchone()[0] == 0:
        c.execute('INSERT INTO admin_users (username, password_hash) VALUES (?, ?)',
                  ('admin', default_pw))
    # Default AI config
    c.execute('SELECT COUNT(*) FROM ai_config')
    if c.fetchone()[0] == 0:
        c.execute('''INSERT INTO ai_config (name, api_key, model, base_url, api_type, is_active)
                     VALUES (?, ?, ?, ?, ?, ?)''',
                  ('DeepSeek V3', 'your-api-key-here',
                   'deepseek-chat', 'https://api.deepseek.com', 'openai', 1))
    conn.commit()
    conn.close()

def get_active_ai_config():
    conn = get_db()
    row = conn.execute('SELECT * FROM ai_config WHERE is_active = 1 LIMIT 1').fetchone()
    conn.close()
    return dict(row) if row else None


def get_site_config(key, default=''):
    conn = get_db()
    row = conn.execute('SELECT value FROM site_config WHERE key = ?', (key,)).fetchone()
    conn.close()
    return row['value'] if row else default


def set_site_config(key, value):
    conn = get_db()
    conn.execute('INSERT OR REPLACE INTO site_config (key, value) VALUES (?, ?)', (key, value))
    conn.commit()
    conn.close()

# ---------- Auth ----------
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated

# ---------- Routes ----------
@app.route('/')
def index():
    return render_template('editor.html')


@app.route('/api/site-config')
def api_site_config():
    """Return site configuration for frontend (drawio_url etc.)."""
    return jsonify({
        'drawio_url': get_site_config('drawio_url', 'https://embed.diagrams.net')
    })

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')
        pw_hash = hashlib.sha256(password.encode()).hexdigest()
        conn = get_db()
        user = conn.execute(
            'SELECT * FROM admin_users WHERE username = ? AND password_hash = ?',
            (username, pw_hash)
        ).fetchone()
        conn.close()
        if user:
            session['admin_logged_in'] = True
            session['admin_user'] = username
            return redirect(url_for('admin_panel'))
        return render_template('admin_login.html', error='用户名或密码错误')
    return render_template('admin_login.html')

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_logged_in', None)
    session.pop('admin_user', None)
    return redirect(url_for('admin_login'))

@app.route('/admin')
@admin_required
def admin_panel():
    conn = get_db()
    configs = conn.execute('SELECT * FROM ai_config ORDER BY id').fetchall()
    conn.close()
    drawio_url = get_site_config('drawio_url', 'https://embed.diagrams.net')
    return render_template('admin.html', configs=[dict(c) for c in configs], drawio_url=drawio_url)

@app.route('/admin/config/add', methods=['POST'])
@admin_required
def add_config():
    data = request.form
    conn = get_db()
    conn.execute(
        'INSERT INTO ai_config (name, api_key, model, base_url, api_type, is_active) VALUES (?,?,?,?,?,?)',
        (data['name'], data['api_key'], data['model'], data['base_url'], data['api_type'], 0)
    )
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))

@app.route('/admin/config/activate/<int:config_id>', methods=['POST'])
@admin_required
def activate_config(config_id):
    conn = get_db()
    conn.execute('UPDATE ai_config SET is_active = 0')
    conn.execute('UPDATE ai_config SET is_active = 1 WHERE id = ?', (config_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))

@app.route('/admin/config/delete/<int:config_id>', methods=['POST'])
@admin_required
def delete_config(config_id):
    conn = get_db()
    conn.execute('DELETE FROM ai_config WHERE id = ?', (config_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/admin/site-config', methods=['POST'])
@admin_required
def save_site_config():
    drawio_url = request.form.get('drawio_url', '').strip()
    if drawio_url:
        set_site_config('drawio_url', drawio_url)
    return redirect(url_for('admin_panel'))

@app.route('/admin/password', methods=['POST'])
@admin_required
def change_password():
    old_pw = request.form.get('old_password', '')
    new_pw = request.form.get('new_password', '')
    old_hash = hashlib.sha256(old_pw.encode()).hexdigest()
    new_hash = hashlib.sha256(new_pw.encode()).hexdigest()
    conn = get_db()
    user = conn.execute(
        'SELECT * FROM admin_users WHERE username = ? AND password_hash = ?',
        (session['admin_user'], old_hash)
    ).fetchone()
    if user:
        conn.execute(
            'UPDATE admin_users SET password_hash = ? WHERE username = ?',
            (new_hash, session['admin_user'])
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    conn.close()
    return jsonify({'success': False, 'error': '旧密码错误'}), 400

# ---------- AI Check API ----------

# Fixed section titles — the ONLY allowed titles in the document
ALL_SECTION_TITLES = [
    '一、目的',
    '二、适用范围',
    '三、规范性引用文件',
    '四、术语和定义',
    '五、基本原则',
    '六、职责',
    '七、管理要求',
    '八、特殊处理机制',
    '九、监管与问责',
    '十、附则',
    '十一、附录',
]

SECTION_PROMPTS = {}

SECTION_PROMPTS['一、目的'] = (
    '审核要点：\n'
    '1. 是否明确了制度制定的核心目标，包括：制度针对的管理对象、要解决的核心问题。\n'
    '2. 是否包含量化可评估的指标（如百分比、金额、时限等），避免"加强""提高"等模糊表述。\n'
    '3. 语句结构是否规范，通常以"为规范……，……，特制定本制度。"的格式。\n'
    '示例：为规范公司合同管理行为，降低法律风险，确保合同履约率≥98%，特制定本制度。'
)

SECTION_PROMPTS['二、适用范围'] = (
    '审核要点：\n'
    '1. 是否明确了适用对象（如全公司/特定部门/特定岗位）。\n'
    '2. 是否明确了适用的业务场景（如采购类合同/劳动合同等）。\n'
    '3. 是否明确了排除范围（如"本制度不适用于XX子公司"）。\n'
    '以上三项（适用对象、适用场景、排除范围）缺少任何一项都应明确指出。'
)

SECTION_PROMPTS['三、规范性引用文件'] = (
    '审核要点：\n'
    '1. 本章节为可选项，填写"无"是允许的。\n'
    '2. 如果填写了引用文件，检查格式是否规范——应包含文件编号和名称（如"GB/T XXXX-XXXX《文件名》"）。\n'
    '3. 引用的文件是否与本制度内容相关。'
)

SECTION_PROMPTS['四、术语和定义'] = (
    '审核要点：\n'
    '1. 本章节为可选项，填写"无"是允许的。\n'
    '2. 如果有术语定义，检查每个术语是否清晰、无歧义。\n'
    '3. 涉及数值边界的术语是否包含了量化标准。\n'
    '示例："重大合同：指单笔金额≥100万元或涉及公司核心技术的合同。"'
)

SECTION_PROMPTS['五、基本原则'] = (
    '审核要点：\n'
    '1. 本章节为可选项，填写"无"是允许的。\n'
    '2. 如果填写了内容，是否列出了3-5条核心管理原则。\n'
    '3. 每条原则是否有名称和具体说明，是否体现了企业管理的价值观或方法论。\n'
    '示例：\n'
    '1．合法性原则：所有合同条款必须符合《民法典》规定；\n'
    '2．风险优先原则：金额每增加50万元，审批层级提高一级。'
)

SECTION_PROMPTS['六、职责'] = (
    '审核要点：\n'
    '1. 是否为每个相关部门或岗位明确了具体职责内容。\n'
    '2. 是否明确了权限边界（如"否决权仅限法律条款"）。\n'
    '3. 各部门/岗位的职责是否清晰、无交叉重叠、无遗漏。\n'
    '4. 职责描述是否使用了可操作的动词（如"审查""审批""备案""监督"等）。'
)

SECTION_PROMPTS['七、管理要求'] = (
    '审核要点：\n'
    '1. 是否使用了分条款的形式（如7.1、7.2……）列出管理规则。\n'
    '2. 是否区分了禁止性要求（"禁止……"）、强制性要求（"必须……"）和指引性要求（"应当……"）。\n'
    '3. 条款是否具体可执行，是否包含量化要素（数字、金额、时限、比例等）。\n'
    '4. 条款之间是否存在矛盾或重复。\n'
    '示例：\n'
    '禁止任何部门未经招标签订≥20万元的采购合同；\n'
    '所有合同必须使用公司标准模板，修改条款需法务部备案。'
)

SECTION_PROMPTS['八、特殊处理机制'] = (
    '审核要点：\n'
    '1. 本章节为可选项，填写"无"是允许的。\n'
    '2. 如果填写了内容，是否明确了特殊情况的定义（什么情况算"特殊"）。\n'
    '3. 是否明确了处理方式和申请路径（审批人、时限要求等）。\n'
    '示例："因紧急生产需要可先执行后补签合同，但需在3个工作日内提交CEO特批说明。"'
)

SECTION_PROMPTS['九、监管与问责'] = (
    '审核要点：\n'
    '1. 监督机制是否包含两种：\n'
    '   - 定期检查（如"每季度审计部抽查20%合同"）；\n'
    '   - 触发检查（如"单笔合同违约金额≥10万元启动专项审计"）。\n'
    '2. 惩罚措施是否将违规行为与处罚措施对应（重大制度必填，一般制度若没有则填"无"）。\n'
    '3. 惩罚措施是否引用了相关制度条款（如《员工奖惩制度》具体条款）。\n'
    '示例：私自修改合同条款者，按造成损失的200%追偿；未按时归档合同，扣减责任人当月绩效10%。\n'
    '注意：重大制度与一般制度的区分详见《制度新增&修订&废除流程说明书》4.3。'
)

SECTION_PROMPTS['十、附则'] = (
    '审核要点：\n'
    '1. 是否包含生效日期说明（如"本办法自会签发布之日起生效"）。\n'
    '2. 是否设置了试行期（重大制度通常3-6个月，一般制度1个月），并说明根据反馈调整。\n'
    '3. 是否明确了解释权归属（如"本办法的最终解释权归XX部门"）。\n'
    '以上三项缺少任何一项都应明确指出。'
)

SECTION_PROMPTS['十一、附录'] = (
    '审核要点：\n'
    '1. 本章节为可选项，填写"无"是允许的。\n'
    '2. 如果填写了内容，是否列出了必要的附件清单。\n'
    '3. 常见附件包括：审批流程图（建议必附）、标准模板（如合同/表单）、记录表（如检查台账）等。\n'
    '4. 附件列表格式是否规范（如使用编号列出）。'
)

@app.route('/api/ai-check', methods=['POST'])
def ai_check():
    data = request.json
    section_title = data.get('title', '')
    content = data.get('content', '')
    doc_name = data.get('doc_name', '').strip()

    if not content.strip():
        return jsonify({'error': '内容为空，无法检查'}), 400

    config = get_active_ai_config()
    if not config:
        return jsonify({'error': '未配置AI模型，请联系管理员'}), 500

    section_hint = SECTION_PROMPTS.get(section_title, '检查该部分内容是否完善、逻辑清晰。')

    all_titles_str = '、'.join(ALL_SECTION_TITLES)

    # Build document name context for the prompt
    if doc_name:
        doc_name_block = (
            f'## 制度名称：{doc_name}\n\n'
            f'用户正在编写的制度名称为「{doc_name}」。审核时请结合此名称判断：\n'
            f'- 本章节的内容是否与「{doc_name}」的主题相关、一致。\n'
            f'- 用语和术语是否贴合「{doc_name}」所涉及的业务领域。\n'
            f'- 如果内容明显偏离「{doc_name}」的主题，应指出并给出修正建议。\n\n'
        )
    else:
        doc_name_block = ''

    system_prompt = (
        '你是一位专业的企业制度审核专家。你需要对用户提交的制度文档中的某个章节进行审核。\n\n'

        f'{doc_name_block}'

        '## 制度文档固定结构\n\n'
        '本制度文档有且仅有以下11个章节标题，标题是固定不可更改的：\n'
        f'{all_titles_str}\n\n'

        '**重要约束——关于标题的规则：**\n'
        '- 这11个标题是唯一合法的章节标题，不可以建议用户新增、删除、合并或重命名任何标题。\n'
        '- 如果你发现用户填写的内容与当前章节标题不匹配（例如把"职责"内容写到了"管理要求"章节），'
        '你应该建议用户将该内容移至上述11个标题中对应的正确章节，而不是建议更改标题。\n'
        '- 你只可以建议"此内容更适合放在「六、职责」章节"这类引用已有标题的建议，'
        '绝不可以建议"建议将标题改为XX"或"建议增加一个XX章节"。\n\n'

        f'## 当前审核章节：{section_title}\n\n'
        f'{section_hint}\n\n'

        '## 审核维度\n\n'
        f'1. **内容匹配度**：内容是否属于"{section_title}"这个章节，是否有内容错放到了本章节（应建议移至正确的章节）。\n'
        '2. **完整性**：是否涵盖了该章节应有的要素，缺少哪些必要内容。\n'
        '3. **规范性**：表述是否规范、专业、无歧义，用词是否准确。\n'
        '4. **可操作性**：条款是否具体、可执行、可量化（包含数字、时限、金额等）。\n'
        '5. **与其他章节的衔接**：内容是否与其他章节存在重复或矛盾。\n'
        '6. **与制度主题的一致性**：内容是否紧扣制度名称所涉及的业务领域，用语和示例是否贴合该制度主题。\n\n'

        '## 输出格式\n\n'
        '请给出：\n'
        '- **总体评分**（满分10分）\n'
        '- **存在的问题**（逐条列出，如有）\n'
        '- **具体修改建议**（逐条列出，如有）\n\n'
        '回复请使用中文，简洁明了，直接指出问题和建议。不要输出与审核无关的寒暄内容。'
    )

    # Strip HTML tags for AI analysis
    soup = BeautifulSoup(content, 'html.parser')
    plain_text = soup.get_text(separator='\n')
    user_message = f'请审核以下内容：\n\n{plain_text}'

    try:
        result = _call_ai(config, system_prompt, user_message)
        return jsonify({'result': result})
    except Exception as e:
        return jsonify({'error': f'AI服务调用失败：{str(e)}'}), 500


def _call_ai(config, system_prompt, user_message):
    """Unified AI call supporting OpenAI-compatible and Anthropic protocols."""
    api_type = config.get('api_type', 'openai')

    if api_type == 'anthropic':
        # Anthropic Messages API:
        #   system prompt is a separate parameter
        #   response: message.content[0].text
        client = anthropic.Anthropic(
            api_key=config['api_key'],
            base_url=config['base_url'] or None
        )
        message = client.messages.create(
            model=config['model'],
            max_tokens=1500,
            system=system_prompt,
            messages=[
                {'role': 'user', 'content': user_message}
            ],
            temperature=0.3
        )
        return message.content[0].text
    else:
        # OpenAI-compatible API (DeepSeek, OpenAI, Moonshot, etc.):
        #   system prompt in messages array
        #   response: response.choices[0].message.content
        client = OpenAI(
            api_key=config['api_key'],
            base_url=config['base_url']
        )
        response = client.chat.completions.create(
            model=config['model'],
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_message}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        return response.choices[0].message.content

# ---------- DOCX Export ----------
@app.route('/api/export', methods=['POST'])
def export_docx():
    data = request.json
    sections = data.get('sections', [])

    from docx.shared import RGBColor
    from docx.shared import Pt as PtSize
    from docx.enum.text import WD_LINE_SPACING

    # Check if template exists
    template_path = os.path.join(_app_dir(), 'template.docx')
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # --- Save template tables' column widths before any style changes ---
    _saved_table_widths = []
    for tbl in doc.tables:
        col_widths = []
        for col in tbl.columns:
            col_widths.append(col.width)
        _saved_table_widths.append((tbl, col_widths))
    # Also save header/footer table widths
    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer,
                        section.even_page_header, section.even_page_footer,
                        section.first_page_header, section.first_page_footer):
            try:
                if hdr_ftr and not hdr_ftr.is_linked_to_previous:
                    for tbl in hdr_ftr._element.findall(qn('w:tbl')):
                        from docx.table import Table
                        t = Table(tbl, doc)
                        col_widths = [c.width for c in t.columns]
                        _saved_table_widths.append((t, col_widths))
            except Exception:
                pass

    # Create a custom style for body content (avoid touching Normal which
    # would cascade to header/footer and template tables)
    from docx.enum.style import WD_STYLE_TYPE
    try:
        body_style = doc.styles['ZhiduBody']
    except KeyError:
        body_style = doc.styles.add_style('ZhiduBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.base_style = doc.styles['Normal']
    body_font = body_style.font
    body_font.name = '宋体'
    body_font.size = Pt(12)  # 小四号 = 12pt
    body_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 正文：段前1行 段后1行 1.5倍行距
    body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _set_spacing_in_lines(body_style.element, before_lines=100, after_lines=100)

    # Document title: 制度名称，黑体24号加粗，单倍行距，居中
    doc_name = data.get('doc_name', '').strip() or 'XXX管理办法'
    title_para = doc.add_paragraph(style='ZhiduBody')
    title_run = title_para.add_run(doc_name)
    title_run.font.name = '黑体'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for sec in sections:
        title = sec.get('title', '')
        content_html = sec.get('content', '')

        # Add heading as a normal paragraph (avoids the black dot from Heading styles)
        heading_para = doc.add_paragraph(style='ZhiduBody')
        heading_run = heading_para.add_run(title)
        heading_run.font.name = '黑体'
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 一级标题：段前1行 段后2行 单倍行距
        heading_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _set_spacing_in_lines(heading_para._element, before_lines=100, after_lines=200)

        # Parse HTML content and add as paragraphs
        if content_html.strip():
            content_html = _clean_word_html(content_html)
            # Log HTML with images for debugging size/alignment issues
            if '<img' in content_html:
                app.logger.info(f'Export HTML with images in [{title}]: {content_html[:2000]}')
            soup = BeautifulSoup(content_html, 'html.parser')
            _parse_html_to_docx(doc, soup)

    # --- Restore saved template table widths ---
    for tbl, col_widths in _saved_table_widths:
        try:
            tbl.allow_autofit = False
            for i, w in enumerate(col_widths):
                if w and i < len(tbl.columns):
                    tbl.columns[i].width = w
                    for row in tbl.rows:
                        if i < len(row.cells):
                            row.cells[i].width = w
        except Exception:
            pass

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='制度文档.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def _clean_word_html(html):
    """Remove Microsoft Word / Office HTML artifacts."""
    # 1. Conditional comments: <!--[if ...]>...<![endif]-->
    html = re.sub(r'<!--\[if[\s\S]*?<!\[endif\]-->', '', html)
    # 2. Stray [if !supportLists] / [endif] (without comment wrapper)
    html = re.sub(r'\[if\s+!support\w+\]', '', html)
    html = re.sub(r'\[endif\]', '', html)
    # 3. XML declarations <?xml ...?>
    html = re.sub(r'<\?xml[\s\S]*?\?>', '', html)
    # 4. <style> blocks (Word embeds massive mso style blocks)
    html = re.sub(r'<style[\s\S]*?</style>', '', html, flags=re.IGNORECASE)
    # 5. Namespace tags: <o:p>, </o:p>, <v:shape>, <w:wrap>, <st1:*>, <m:*> etc.
    html = re.sub(r'</?\w+:[^>]*>', '', html)
    # 6. xmlns:* attributes
    html = re.sub(r'\s*xmlns:\w+="[^"]*"', '', html)
    # 7. class="Mso..." attributes
    html = re.sub(r'\s*class="Mso[^"]*"', '', html)
    # 8. mso-* CSS properties inside style attributes
    html = re.sub(r'mso-[a-z\-]+\s*:[^;"]*;?', '', html, flags=re.IGNORECASE)
    # 9. Empty style="" left after mso removal
    html = re.sub(r'\s*style="\s*"', '', html)
    # 10. <font> tags (keep content)
    html = re.sub(r'</?font[^>]*>', '', html, flags=re.IGNORECASE)
    # 11. <span style="mso-spacerun:yes"> &nbsp; </span>
    html = re.sub(r'<span[^>]*mso-spacerun[^>]*>[\s\u00a0]*</span>', '', html, flags=re.IGNORECASE)
    # 12. Empty <span></span>
    html = re.sub(r'<span></span>', '', html)
    # 13. lang / xml:lang attributes
    html = re.sub(r'\s*(xml:)?lang="[^"]*"', '', html)
    # 14. Empty class="" attributes
    html = re.sub(r'\s*class="\s*"', '', html)
    return html


def _is_block_element(tag_name):
    """Check if an HTML tag is a block-level element."""
    return tag_name in (
        'p', 'div', 'table', 'ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'blockquote', 'pre', 'figure', 'section', 'article', 'main', 'header',
        'footer', 'nav', 'aside', 'hr', 'img',
    )


def _parse_html_to_docx(doc, soup):
    """Parse HTML content and add to docx document."""
    from docx.shared import RGBColor, Inches

    _style = 'ZhiduBody'
    for element in soup.children:
        if element.name == 'img':
            _add_image_to_docx(doc, element)
        elif element.name == 'table':
            _add_table_to_docx(doc, element)
        elif element.name == 'p':
            _add_paragraph_to_docx(doc, element)
        elif element.name is None:
            text = str(element).strip()
            if text:
                para = doc.add_paragraph(style=_style)
                _set_run_font(para.add_run(text))
        elif element.name in ('ul', 'ol'):
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style=_style)
                _add_formatted_runs(para, li, li.get_text())
                try:
                    if element.name == 'ul':
                        para.style = doc.styles['List Bullet']
                    else:
                        para.style = doc.styles['List Number']
                except KeyError:
                    # Style doesn't exist in template, use manual formatting
                    if element.name == 'ul':
                        para.paragraph_format.left_indent = Pt(24)
                        para.paragraph_format.first_line_indent = Pt(-12)
                        para.add_run('• ').insert_paragraph_before()
                        para.text = '• ' + para.text
                    else:
                        para.paragraph_format.left_indent = Pt(24)
                        para.paragraph_format.first_line_indent = Pt(-12)
        elif element.name in ('h1','h2','h3','h4','h5','h6'):
            para = doc.add_paragraph(style=_style)
            run = para.add_run(element.get_text())
            run.bold = True
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif element.name:
            # Container elements (div, figure, section, etc.):
            # check if they contain block-level children and recurse
            has_block = any(
                hasattr(c, 'name') and c.name and _is_block_element(c.name)
                for c in element.children
            )
            if has_block:
                _parse_html_to_docx(doc, element)
            else:
                # Treat as an inline/text container
                text = element.get_text()
                if text.strip():
                    para = doc.add_paragraph(style=_style)
                    _add_formatted_runs(para, element, text)
                    _apply_paragraph_format(para, element)
        elif isinstance(element, str) and element.strip():
            para = doc.add_paragraph(style=_style)
            _set_run_font(para.add_run(element.strip()))


def _add_paragraph_to_docx(doc, element):
    """Add a <p> element to docx, handling images and formatting."""
    _style = 'ZhiduBody'
    imgs = element.find_all('img')
    if imgs:
        for child in element.children:
            if hasattr(child, 'name') and child.name == 'img':
                _add_image_to_docx(doc, child)
            elif isinstance(child, str) and child.strip():
                para = doc.add_paragraph(style=_style)
                _set_run_font(para.add_run(child.strip()))
                _apply_paragraph_format(para, element)
            elif hasattr(child, 'name') and child.name:
                text = child.get_text()
                if text.strip():
                    para = doc.add_paragraph(style=_style)
                    _add_formatted_runs(para, child, text)
                    _apply_paragraph_format(para, element)
    else:
        text = element.get_text()
        if text.strip():
            para = doc.add_paragraph(style=_style)
            _add_formatted_runs(para, element, text)
            _apply_paragraph_format(para, element)


def _set_spacing_in_lines(element, before_lines=None, after_lines=None):
    """Set paragraph before/after spacing in 'lines' unit via XML.
    element: paragraph._element (CT_P) or style.element (CT_Style)
    Values: 100 = 1 line, 200 = 2 lines, etc.
    python-docx has no high-level API for w:beforeLines/w:afterLines.
    """
    from docx.oxml import OxmlElement
    pPr = element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before_lines is not None:
        spacing.set(qn('w:beforeLines'), str(before_lines))
        if qn('w:before') in spacing.attrib:
            del spacing.attrib[qn('w:before')]
    if after_lines is not None:
        spacing.set(qn('w:afterLines'), str(after_lines))
        if qn('w:after') in spacing.attrib:
            del spacing.attrib[qn('w:after')]


def _apply_paragraph_format(para, element):
    """Apply text-align and text-indent from HTML style to docx paragraph."""
    if not element or not element.get('style'):
        return
    style_str = element.get('style', '')
    # Horizontal alignment
    m = re.search(r'text-align:\s*(left|center|right|justify)', style_str, re.IGNORECASE)
    if m:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = align_map.get(m.group(1).lower())
    # First-line indent
    if 'text-indent' in style_str:
        m = re.search(r'text-indent:\s*([\d.]+)\s*em', style_str)
        if m:
            em_val = float(m.group(1))
            para.paragraph_format.first_line_indent = Pt(int(em_val * 12))


def _parse_img_size(img_tag):
    """Extract width and height in EMU from <img> style or attributes.
    Returns (width_emu, height_emu); either may be None if not specified.
    1 CSS pixel = 9525 EMU.
    """
    from docx.shared import Emu
    width_px = None
    height_px = None

    # 1) Try style attribute: style="width: 300px; height: 200px"
    style = img_tag.get('style', '')
    if style:
        w_match = re.search(r'(?<![a-z-])width:\s*([\d.]+)\s*px', style)
        h_match = re.search(r'(?<![a-z-])height:\s*([\d.]+)\s*px', style)
        if w_match:
            width_px = float(w_match.group(1))
        if h_match:
            height_px = float(h_match.group(1))

    # 2) Fall back to HTML attributes: width="300" height="200"
    if width_px is None and img_tag.get('width'):
        try:
            val = str(img_tag['width']).strip().replace('px', '')
            if val and not val.endswith('%'):
                width_px = float(val)
        except (ValueError, AttributeError):
            pass
    if height_px is None and img_tag.get('height'):
        try:
            val = str(img_tag['height']).strip().replace('px', '')
            if val and not val.endswith('%'):
                height_px = float(val)
        except (ValueError, AttributeError):
            pass

    app.logger.debug(f'_parse_img_size: width_px={width_px}, height_px={height_px}, '
                     f'style={style!r}, w_attr={img_tag.get("width")}, h_attr={img_tag.get("height")}')

    w_emu = Emu(int(width_px * 9525)) if width_px is not None else None
    h_emu = Emu(int(height_px * 9525)) if height_px is not None else None
    return w_emu, h_emu


def _get_img_alignment(img_tag):
    """Detect image alignment from <img> style and parent element.
    TinyMCE sets alignment on the <img> itself:
      - Center: display: block; margin-left: auto; margin-right: auto
      - Right:  float: right
      - Left:   float: left  (or default — no explicit style)
    Also checks parent <p> text-align as fallback.
    Returns WD_ALIGN_PARAGRAPH value or None (Word default = left).
    """
    style = (img_tag.get('style') or '').lower()

    # Check image's own style first (TinyMCE default image alignment)
    float_m = re.search(r'float\s*:\s*(left|right)', style)
    if float_m:
        if float_m.group(1) == 'right':
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.LEFT
    if re.search(r'margin-left\s*:\s*auto', style) and re.search(r'margin-right\s*:\s*auto', style):
        return WD_ALIGN_PARAGRAPH.CENTER

    # Check parent and ancestor <p>/<div> for text-align
    node = img_tag.parent
    while node and hasattr(node, 'name'):
        if node.name in ('p', 'div', 'td', 'th'):
            ps = (node.get('style') or '').lower()
            m = re.search(r'text-align:\s*(left|center|right|justify)', ps)
            if m:
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                return align_map.get(m.group(1))
            break
        node = node.parent

    return None


def _resolve_img_stream(img_tag):
    """Resolve <img> src to a file path or BytesIO stream. Returns None on failure."""
    import base64
    src = img_tag.get('src', '')
    if not src:
        return None

    if src.startswith('/uploads/'):
        img_path = os.path.join(UPLOAD_FOLDER, src.replace('/uploads/', '', 1))
        if os.path.exists(img_path):
            return img_path
    elif '/uploads/' in src and not src.startswith('blob:'):
        filename = src.split('/uploads/')[-1].split('?')[0]
        img_path = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.exists(img_path):
            return img_path
    elif src.startswith('data:image/'):
        match = re.match(r'data:image/[^;]+;base64,(.*)', src, re.DOTALL)
        if match:
            return io.BytesIO(base64.b64decode(match.group(1)))
    return None


def _apply_img_size(pic, target_w, target_h, max_width):
    """Apply parsed dimensions to a picture, enforcing max_width."""
    if target_w is not None:
        if target_h is not None:
            pic.width = int(target_w)
            pic.height = int(target_h)
        else:
            ratio = target_w / pic.width if pic.width else 1
            pic.height = int(pic.height * ratio)
            pic.width = int(target_w)
    elif target_h is not None:
        ratio = target_h / pic.height if pic.height else 1
        pic.width = int(pic.width * ratio)
        pic.height = int(target_h)
    # Enforce max width
    if pic.width > max_width:
        ratio = max_width / pic.width
        pic.width = int(max_width)
        pic.height = int(pic.height * ratio)


def _add_image_to_docx(doc, img_tag):
    """Add an image from an <img> tag to the docx document."""
    img_stream = _resolve_img_stream(img_tag)
    if img_stream is None:
        src = img_tag.get('src', '')
        app.logger.warning(f'Image not resolved, src={src[:100]}')
        return

    try:
        target_w, target_h = _parse_img_size(img_tag)
        pic = doc.add_picture(img_stream)
        section = doc.sections[-1]
        max_width = section.page_width - section.left_margin - section.right_margin
        _apply_img_size(pic, target_w, target_h, max_width)

        para = doc.paragraphs[-1]
        try:
            para.style = doc.styles['ZhiduBody']
        except KeyError:
            pass
        align = _get_img_alignment(img_tag)
        if align is not None:
            para.alignment = align
    except Exception as e:
        app.logger.error(f'Failed to add image to docx: {e}')

def _add_formatted_runs(para, element, fallback_text):
    """Add runs with formatting from HTML element."""
    from docx.shared import RGBColor
    if element is None:
        _set_run_font(para.add_run(fallback_text))
        return

    for child in element.children if element else []:
        if isinstance(child, str):
            if child.strip():
                _set_run_font(para.add_run(child))
        elif child.name == 'strong' or child.name == 'b':
            run = para.add_run(child.get_text())
            run.bold = True
            _set_run_font(run)
        elif child.name == 'em' or child.name == 'i':
            run = para.add_run(child.get_text())
            run.italic = True
            _set_run_font(run)
        elif child.name == 'u':
            run = para.add_run(child.get_text())
            run.underline = True
            _set_run_font(run)
        elif child.name == 'span':
            # Handle span with inline styles (color, background, strikethrough)
            run = para.add_run(child.get_text())
            _set_run_font(run)
            style_str = child.get('style', '')
            # Text color
            color_match = re.search(r'color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', style_str)
            if not color_match:
                color_match = re.search(r'color:\s*#([0-9a-fA-F]{6})', style_str)
                if color_match:
                    hex_color = color_match.group(1)
                    run.font.color.rgb = RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
            else:
                run.font.color.rgb = RGBColor(int(color_match.group(1)), int(color_match.group(2)), int(color_match.group(3)))
            # Background color (highlight)
            bg_match = re.search(r'background-color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', style_str)
            if not bg_match:
                bg_match = re.search(r'background-color:\s*#([0-9a-fA-F]{6})', style_str)
                if bg_match:
                    hex_color = bg_match.group(1)
                    run.font.highlight_color = _rgb_to_highlight(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
            else:
                run.font.highlight_color = _rgb_to_highlight(int(bg_match.group(1)), int(bg_match.group(2)), int(bg_match.group(3)))
            # Text decoration: line-through
            if 'text-decoration: line-through' in style_str or 'text-decoration:line-through' in style_str:
                run.font.strike = True
        elif child.name == 's' or child.name == 'strike' or child.name == 'del':
            run = para.add_run(child.get_text())
            run.font.strike = True
            _set_run_font(run)
        else:
            _set_run_font(para.add_run(child.get_text()))

def _rgb_to_highlight(r, g, b):
    """Map RGB to closest WD_COLOR_INDEX highlight color."""
    from docx.enum.text import WD_COLOR_INDEX
    # Simple mapping to common highlight colors
    if r > 200 and g > 200 and b < 100:
        return WD_COLOR_INDEX.YELLOW
    elif r > 200 and g < 100 and b < 100:
        return WD_COLOR_INDEX.RED
    elif r < 100 and g > 200 and b < 100:
        return WD_COLOR_INDEX.GREEN
    elif r < 100 and g < 100 and b > 200:
        return WD_COLOR_INDEX.BLUE
    elif r > 200 and g > 150 and b < 100:
        return WD_COLOR_INDEX.DARK_YELLOW
    return None  # No highlight

def _set_run_font(run):
    """Set font for a run: Chinese = 宋体, English = Times New Roman, size = 12pt."""
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def _parse_width(style_str, width_attr):
    """Parse width from CSS style string or HTML width attribute.
    Returns (value, unit) tuple e.g. (30, '%') or (200, 'px'), or None.
    """
    import re
    # Try CSS style first: width: 30% or width: 200px
    if style_str:
        m = re.search(r'width\s*:\s*([\d.]+)\s*(px|%)', style_str)
        if m:
            return (float(m.group(1)), m.group(2))
    # Fallback to HTML width attribute: "30%" or "200"
    if width_attr:
        width_attr = str(width_attr).strip()
        m = re.match(r'^([\d.]+)\s*(%|px)?$', width_attr)
        if m:
            unit = m.group(2) or 'px'
            return (float(m.group(1)), unit)
    return None


def _add_table_to_docx(doc, table_elem):
    """Convert HTML table to docx table, preserving column widths and merged cells."""
    from docx.shared import Inches, Emu
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import base64
    import re

    def set_cell_border(cell, **kwargs):
        """Set cell border. kwargs: top, bottom, left, right, insideH, insideV"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            if edge in kwargs:
                edge_data = kwargs[edge]
                edge_el = OxmlElement(f'w:{edge}')
                for key, value in edge_data.items():
                    edge_el.set(qn(f'w:{key}'), str(value))
                tcBorders.append(edge_el)
        tcPr.append(tcBorders)

    rows_html = table_elem.find_all('tr')
    if not rows_html:
        return
    num_rows = len(rows_html)

    # --- Phase 1: Build grid, accounting for colspan & rowspan ---
    cell_map = {}   # (row, col) -> {'cell': element, 'colspan': int, 'rowspan': int}
    occupied = set()  # positions taken by a span (not the origin cell)

    for i, row_elem in enumerate(rows_html):
        html_cells = row_elem.find_all(['td', 'th'])
        col = 0
        for html_cell in html_cells:
            while (i, col) in occupied or (i, col) in cell_map:
                col += 1
            colspan = int(html_cell.get('colspan', 1) or 1)
            rowspan = int(html_cell.get('rowspan', 1) or 1)
            cell_map[(i, col)] = {
                'cell': html_cell, 'colspan': colspan, 'rowspan': rowspan,
            }
            for dr in range(rowspan):
                for dc in range(colspan):
                    if dr > 0 or dc > 0:
                        occupied.add((i + dr, col + dc))
            col += colspan

    all_positions = set(cell_map.keys()) | occupied
    if not all_positions:
        return
    actual_cols = max(c + 1 for _, c in all_positions)

    # --- Phase 2: Create docx table ---
    table = doc.add_table(rows=num_rows, cols=actual_cols)

    style_applied = False
    try:
        table.style = 'Table Grid'
        style_applied = True
    except KeyError:
        pass

    # --- Phase 3: Merge cells ---
    for (r, c), info in cell_map.items():
        cs, rs = info['colspan'], info['rowspan']
        if cs > 1 or rs > 1:
            table.cell(r, c).merge(table.cell(r + rs - 1, c + cs - 1))

    # Manual borders after merge so they apply to final cell shapes
    if not style_applied:
        border_config = {"sz": 4, "val": "single", "color": "000000"}
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top=border_config, bottom=border_config,
                                left=border_config, right=border_config)

    # --- Phase 4: Table alignment ---
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table_style_str = table_elem.get('style', '')
    table_align_attr = table_elem.get('align', '')

    if 'margin-left: auto' in table_style_str and 'margin-right: auto' in table_style_str:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    elif 'margin-left: auto' in table_style_str:
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    elif table_align_attr:
        align_map = {'left': WD_TABLE_ALIGNMENT.LEFT, 'center': WD_TABLE_ALIGNMENT.CENTER, 'right': WD_TABLE_ALIGNMENT.RIGHT}
        table.alignment = align_map.get(table_align_attr.lower(), WD_TABLE_ALIGNMENT.LEFT)
    elif 'float: right' in table_style_str or 'float:right' in table_style_str:
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # --- Phase 5: Resolve column widths ---
    total_width_inches = 6.0
    col_widths = None

    colgroup = table_elem.find('colgroup')
    if colgroup:
        cols_tags = colgroup.find_all('col')
        if cols_tags:
            col_widths = [_parse_width(ct.get('style', ''), ct.get('width', ''))
                          for ct in cols_tags]

    if not col_widths:
        first_row_cells = rows_html[0].find_all(['td', 'th'])
        if first_row_cells:
            widths = []
            for fc in first_row_cells:
                w = _parse_width(fc.get('style', ''), fc.get('width', ''))
                cs = int(fc.get('colspan', 1) or 1)
                if cs > 1 and w:
                    per_col = (w[0] / cs, w[1])
                    widths.extend([per_col] * cs)
                else:
                    widths.append(w)
            if any(w is not None for w in widths):
                col_widths = widths

    if col_widths:
        while len(col_widths) < actual_cols:
            col_widths.append(None)
        col_widths = col_widths[:actual_cols]

        resolved = [None] * actual_cols
        has_pct = any(isinstance(w, tuple) and w[1] == '%' for w in col_widths if w)
        has_px = any(isinstance(w, tuple) and w[1] == 'px' for w in col_widths if w)

        if has_pct:
            for j, w in enumerate(col_widths):
                if w and w[1] == '%':
                    resolved[j] = total_width_inches * w[0] / 100.0
        elif has_px:
            total_px = sum(w[0] for w in col_widths if w)
            if total_px > 0:
                for j, w in enumerate(col_widths):
                    if w and w[1] == 'px':
                        resolved[j] = total_width_inches * w[0] / total_px

        table.autofit = False
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        for j, w_inches in enumerate(resolved):
            if w_inches:
                table.columns[j].width = Inches(w_inches)
                for row in table.rows:
                    row.cells[j].width = Inches(w_inches)

    # --- Phase 6: Fill cell content ---
    for (r, c), info in cell_map.items():
        html_cell = info['cell']
        docx_cell = table.cell(r, c)
        docx_cell.text = ''

        cell_style = html_cell.get('style', '')
        va_match = re.search(r'vertical-align:\s*(top|middle|bottom)', cell_style, re.IGNORECASE)
        if va_match:
            va_map = {'top': WD_ALIGN_VERTICAL.TOP, 'middle': WD_ALIGN_VERTICAL.CENTER, 'bottom': WD_ALIGN_VERTICAL.BOTTOM}
            docx_cell.vertical_alignment = va_map.get(va_match.group(1).lower(), WD_ALIGN_VERTICAL.CENTER)
        else:
            docx_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Get cell width for constraining images
        cell_width = docx_cell.width
        _fill_cell(doc, docx_cell, html_cell, cell_width)


def _fill_cell(doc, docx_cell, html_cell, cell_width=None):
    """Fill a docx table cell with HTML content including images."""
    from docx.shared import Inches
    import base64
    import re

    # Resolve cell-level horizontal alignment (from <td style="text-align:...">)
    cell_align = None
    cell_style = html_cell.get('style', '')
    align_m = re.search(r'text-align:\s*(left|center|right|justify)', cell_style, re.IGNORECASE)
    if align_m:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        cell_align = align_map.get(align_m.group(1).lower())

    first = True
    for child in html_cell.children:
        if isinstance(child, str):
            text = child.strip()
            if text:
                para = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
                first = False
                _set_run_font(para.add_run(text))
                if cell_align is not None:
                    para.alignment = cell_align
        elif child.name == 'img':
            para = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
            first = False
            _add_image_to_cell(para, child, cell_width)
            if cell_align is not None:
                para.alignment = cell_align
        elif child.name == 'p':
            para = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
            first = False
            # <p> own style takes priority; fall back to cell-level alignment
            _apply_paragraph_format(para, child)
            if para.alignment is None and cell_align is not None:
                para.alignment = cell_align
            # Check for images inside <p>
            imgs = child.find_all('img')
            if imgs:
                for sub in child.children:
                    if hasattr(sub, 'name') and sub.name == 'img':
                        _add_image_to_cell(para, sub, cell_width)
                    elif isinstance(sub, str) and sub.strip():
                        _set_run_font(para.add_run(sub.strip()))
                    elif hasattr(sub, 'name') and sub.name:
                        text = sub.get_text()
                        if text.strip():
                            _set_run_font(para.add_run(text))
            else:
                _add_formatted_runs(para, child, child.get_text())
        elif child.name in ('strong', 'b', 'em', 'i', 'u', 'span'):
            para = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
            first = False
            _add_formatted_runs(para, html_cell, child.get_text())
            if cell_align is not None:
                para.alignment = cell_align
        elif child.name == 'br':
            pass
        elif child.name:
            text = child.get_text(strip=True)
            if text:
                para = docx_cell.paragraphs[0] if first else docx_cell.add_paragraph()
                first = False
                _set_run_font(para.add_run(text))
                if cell_align is not None:
                    para.alignment = cell_align


def _add_image_to_cell(para, img_tag, cell_width=None):
    """Add an image into a table cell paragraph, constrained to cell width."""
    from docx.shared import Inches
    img_stream = _resolve_img_stream(img_tag)
    if img_stream is None:
        return

    try:
        target_w, target_h = _parse_img_size(img_tag)
        run = para.add_run()
        pic = run.add_picture(img_stream)
        # Use cell width as constraint; fall back to Inches(6) if unknown
        max_width = cell_width if cell_width else Inches(6)
        _apply_img_size(pic, target_w, target_h, max_width)
        # Apply alignment from image style
        align = _get_img_alignment(img_tag)
        if align is not None:
            para.alignment = align
    except Exception as e:
        app.logger.error(f'Failed to add image to table cell: {e}')

# ---------- Main ----------
if __name__ == '__main__':
    init_db()
    is_frozen = getattr(sys, 'frozen', False)
    if is_frozen:
        import webbrowser, threading
        threading.Timer(1.5, lambda: webbrowser.open('http://127.0.0.1:5000')).start()
        app.run(host='0.0.0.0', port=5000, debug=False)
    else:
        app.run(debug=True, port=5000)
